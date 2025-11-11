"""
Business Continuity Plan Export Service (BC8).

Handles export of business continuity plans to DOCX and PDF formats.
Implements deterministic content hashing for change tracking.
"""
from __future__ import annotations

import hashlib
import io
import json
from datetime import datetime
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, BaseLoader
try:  # pragma: no cover - import guard for optional dependency
    from weasyprint import HTML
except (ImportError, OSError) as exc:  # pragma: no cover - executed only when missing deps
    HTML = None  # type: ignore[assignment]
    _WEASYPRINT_IMPORT_ERROR = exc
else:
    _WEASYPRINT_IMPORT_ERROR = None

from app.repositories import bc3 as bc_repo
from app.repositories import bcp as bcp_repo
from app.repositories import users as user_repo


# ============================================================================
# Content Hash Generation
# ============================================================================

def compute_content_hash(content: dict[str, Any], metadata: dict[str, Any]) -> str:
    """
    Compute deterministic SHA256 hash of plan content and metadata.
    
    Args:
        content: Plan content JSON
        metadata: Export metadata (title, version, date, author)
        
    Returns:
        SHA256 hex digest string
    """
    # Create a stable JSON representation
    combined = {
        "content": content,
        "metadata": metadata,
    }
    # Use sort_keys to ensure deterministic ordering
    json_bytes = json.dumps(combined, sort_keys=True).encode("utf-8")
    return hashlib.sha256(json_bytes).hexdigest()


# ============================================================================
# DOCX Export
# ============================================================================

async def export_to_docx(
    plan_id: int,
    version_id: Optional[int] = None,
) -> tuple[io.BytesIO, str]:
    """
    Export a business continuity plan to DOCX format.
    
    Args:
        plan_id: ID of the plan to export
        version_id: Specific version ID, or None for active version
        
    Returns:
        Tuple of (BytesIO buffer with DOCX data, content hash)
        
    Raises:
        ValueError: If plan or version not found
    """
    # Fetch plan and version data
    plan = await bc_repo.get_plan_by_id(plan_id)
    if not plan:
        raise ValueError(f"Plan {plan_id} not found")
    
    if version_id:
        version = await bc_repo.get_version_by_id(version_id)
        if not version or version["plan_id"] != plan_id:
            raise ValueError(f"Version {version_id} not found for plan {plan_id}")
    else:
        version = await bc_repo.get_active_version(plan_id)
        if not version:
            raise ValueError(f"No active version found for plan {plan_id}")
    
    # Get template if available
    template = None
    if plan.get("template_id"):
        template = await bc_repo.get_template_by_id(plan["template_id"])
    
    # Get author information
    author = await user_repo.get_user_by_id(version["authored_by_user_id"])
    author_name = author.get("name") if author else "Unknown"
    
    # Prepare metadata
    metadata = {
        "plan_title": plan["title"],
        "version_number": version["version_number"],
        "authored_at": version["authored_at_utc"].isoformat() if version.get("authored_at_utc") else None,
        "author_name": author_name,
        "template_name": template.get("name") if template else None,
    }
    
    # Get plan content
    content_json = version.get("content_json") or {}
    
    # Compute content hash
    content_hash = compute_content_hash(content_json, metadata)
    
    # Create DOCX document
    doc = Document()
    
    # Add title and metadata
    _add_docx_header(doc, metadata)
    
    # Add revision metadata section
    _add_docx_revision_metadata(doc, metadata, version)
    
    # Add plan content sections
    await _add_docx_content(doc, content_json, template, plan_id)
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer, content_hash


def _add_docx_header(doc: Document, metadata: dict[str, Any]) -> None:
    """Add document header with title and metadata."""
    # Title
    title = doc.add_heading(metadata["plan_title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version info
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_para.add_run(f"Version {metadata['version_number']}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add spacing
    doc.add_paragraph()


def _add_docx_revision_metadata(
    doc: Document,
    metadata: dict[str, Any],
    version: dict[str, Any],
) -> None:
    """Add revision metadata section."""
    doc.add_heading("Document Information", level=1)
    
    # Create metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    
    # Set column headers
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    
    # Add metadata rows
    table.rows[1].cells[0].text = "Version"
    table.rows[1].cells[1].text = str(metadata["version_number"])
    
    table.rows[2].cells[0].text = "Author"
    table.rows[2].cells[1].text = metadata["author_name"]
    
    table.rows[3].cells[0].text = "Date"
    authored_date = metadata.get("authored_at", "")
    if authored_date:
        authored_date = datetime.fromisoformat(authored_date).strftime("%Y-%m-%d %H:%M UTC")
    table.rows[3].cells[1].text = authored_date
    
    # Add change note if present
    if version.get("summary_change_note"):
        doc.add_paragraph()
        doc.add_heading("Change Summary", level=2)
        doc.add_paragraph(version["summary_change_note"])
    
    doc.add_page_break()


async def _add_docx_content(
    doc: Document,
    content_json: dict[str, Any],
    template: Optional[dict[str, Any]],
    plan_id: int,
) -> None:
    """Add plan content sections to DOCX."""
    # Get template schema if available
    schema = template.get("schema_json") if template else None
    sections = schema.get("sections", []) if schema else []
    
    # If we have a template schema, use it to structure the content
    if sections:
        for section in sections:
            section_id = section.get("section_id") or section.get("key")
            section_title = section.get("title", section_id)
            section_content = content_json.get(section_id, {})
            
            # Add section heading
            doc.add_heading(section_title, level=1)
            
            # Add section description if present
            if section.get("description"):
                desc_para = doc.add_paragraph(section["description"])
                desc_para.style = "Intense Quote"
            
            # Process fields in the section
            fields = section.get("fields", [])
            for field in fields:
                field_id = field.get("field_id")
                field_label = field.get("label", field_id)
                field_type = field.get("field_type")
                field_value = section_content.get(field_id) if section_content else None
                
                if field_value is None:
                    continue
                
                # Add field based on type
                if field_type == "table":
                    _add_docx_table(doc, field_label, field_value, field.get("columns", []))
                elif field_type in ("text", "rich_text"):
                    doc.add_heading(field_label, level=2)
                    doc.add_paragraph(str(field_value))
                elif field_type in ("date", "datetime"):
                    doc.add_heading(field_label, level=2)
                    doc.add_paragraph(str(field_value))
                else:
                    doc.add_heading(field_label, level=2)
                    doc.add_paragraph(str(field_value))
            
            # Add spacing between sections
            doc.add_paragraph()
    else:
        # No template - just dump content as JSON sections
        for section_key, section_data in content_json.items():
            doc.add_heading(section_key.replace("_", " ").title(), level=1)
            
            if isinstance(section_data, dict):
                for field_key, field_value in section_data.items():
                    doc.add_heading(field_key.replace("_", " ").title(), level=2)
                    doc.add_paragraph(str(field_value))
            else:
                doc.add_paragraph(str(section_data))
    
    # Add additional data tables from database
    await _add_docx_database_tables(doc, plan_id)


def _add_docx_table(
    doc: Document,
    table_title: str,
    table_data: list[dict[str, Any]],
    columns: list[dict[str, Any]],
) -> None:
    """Add a table to the DOCX document."""
    if not table_data:
        return
    
    doc.add_heading(table_title, level=2)
    
    # Determine columns
    if columns:
        col_ids = [col.get("column_id") for col in columns]
        col_labels = [col.get("label", col.get("column_id")) for col in columns]
    else:
        # Infer from first row
        if table_data:
            col_ids = list(table_data[0].keys())
            col_labels = [col.replace("_", " ").title() for col in col_ids]
        else:
            return
    
    # Create table with header row
    table = doc.add_table(rows=1 + len(table_data), cols=len(col_ids))
    table.style = "Light Grid Accent 1"
    
    # Add header row
    header_cells = table.rows[0].cells
    for i, label in enumerate(col_labels):
        header_cells[i].text = label
    
    # Add data rows
    for row_idx, row_data in enumerate(table_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, col_id in enumerate(col_ids):
            value = row_data.get(col_id, "")
            row_cells[col_idx].text = str(value) if value is not None else ""


async def _add_docx_database_tables(doc: Document, plan_id: int) -> None:
    """Add tables from database (BIA, risks, contacts, vendors)."""
    doc.add_page_break()
    doc.add_heading("Appendices", level=1)
    
    # Add risks table
    risks = await bc_repo.list_risks_by_plan(plan_id)
    if risks:
        doc.add_heading("Risk Register", level=2)
        
        risk_table = doc.add_table(rows=1 + len(risks), cols=5)
        risk_table.style = "Light Grid Accent 1"
        
        # Headers
        headers = ["Threat", "Likelihood", "Impact", "Rating", "Mitigation"]
        for i, header in enumerate(headers):
            risk_table.rows[0].cells[i].text = header
        
        # Data rows
        for row_idx, risk in enumerate(risks, start=1):
            risk_table.rows[row_idx].cells[0].text = risk.get("threat", "")
            risk_table.rows[row_idx].cells[1].text = risk.get("likelihood", "")
            risk_table.rows[row_idx].cells[2].text = risk.get("impact", "")
            risk_table.rows[row_idx].cells[3].text = risk.get("rating", "")
            risk_table.rows[row_idx].cells[4].text = risk.get("mitigation", "")
        
        doc.add_paragraph()
    
    # Note: Contacts and vendors would be added here if we had repository functions for them
    # For now, they're likely included in the content_json table fields


# ============================================================================
# PDF Export
# ============================================================================

async def export_to_pdf(
    plan_id: int,
    version_id: Optional[int] = None,
) -> tuple[io.BytesIO, str]:
    """
    Export a business continuity plan to PDF format.
    
    Generates PDF from rendered HTML using Jinja2 templates and WeasyPrint.
    
    Args:
        plan_id: ID of the plan to export
        version_id: Specific version ID, or None for active version
        
    Returns:
        Tuple of (BytesIO buffer with PDF data, content hash)
        
    Raises:
        ValueError: If plan or version not found
    """
    if HTML is None:  # pragma: no cover - depends on system configuration
        raise RuntimeError(
            "PDF export requires WeasyPrint and its native dependencies (libpango and libpangocairo). "
            "Install the system packages documented at "
            "https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation"
        ) from _WEASYPRINT_IMPORT_ERROR

    # Fetch plan and version data
    plan = await bc_repo.get_plan_by_id(plan_id)
    if not plan:
        raise ValueError(f"Plan {plan_id} not found")
    
    if version_id:
        version = await bc_repo.get_version_by_id(version_id)
        if not version or version["plan_id"] != plan_id:
            raise ValueError(f"Version {version_id} not found for plan {plan_id}")
    else:
        version = await bc_repo.get_active_version(plan_id)
        if not version:
            raise ValueError(f"No active version found for plan {plan_id}")
    
    # Get template if available
    template = None
    if plan.get("template_id"):
        template = await bc_repo.get_template_by_id(plan["template_id"])
    
    # Get author information
    author = await user_repo.get_user_by_id(version["authored_by_user_id"])
    author_name = author.get("name") if author else "Unknown"
    
    # Prepare metadata
    metadata = {
        "plan_title": plan["title"],
        "version_number": version["version_number"],
        "authored_at": version["authored_at_utc"].isoformat() if version.get("authored_at_utc") else None,
        "author_name": author_name,
        "template_name": template.get("name") if template else None,
    }
    
    # Get plan content
    content_json = version.get("content_json") or {}
    
    # Compute content hash (same as DOCX for consistency)
    content_hash = compute_content_hash(content_json, metadata)
    
    # Get additional data
    risks = await bc_repo.list_risks_by_plan(plan_id)
    
    # Render HTML from template
    html_content = _render_plan_html(
        plan=plan,
        version=version,
        metadata=metadata,
        content_json=content_json,
        template=template,
        risks=risks,
    )
    
    # Convert HTML to PDF using WeasyPrint
    pdf_buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(pdf_buffer)
    pdf_buffer.seek(0)
    
    return pdf_buffer, content_hash


async def export_bcp_to_pdf(
    plan_id: int,
    event_log_limit: int = 100,
) -> tuple[io.BytesIO, str]:
    """
    Export a Business Continuity Plan to template-faithful PDF format.
    
    Generates comprehensive PDF with all BCP sections in prescribed order:
    1. Plan Overview
    2. Risk Management
    3. Business Impact Analysis
    4. Incident Response  
    5. Recovery
    6. Rehearse/Maintain/Review
    
    Includes footer attribution and configurable event log entries.
    
    Args:
        plan_id: ID of the BCP plan to export
        event_log_limit: Maximum number of event log entries to include (default 100)
        
    Returns:
        Tuple of (BytesIO buffer with PDF data, content hash)
        
    Raises:
        ValueError: If plan not found
        RuntimeError: If WeasyPrint not available
    """
    if HTML is None:  # pragma: no cover - depends on system configuration
        raise RuntimeError(
            "PDF export requires WeasyPrint and its native dependencies (libpango and libpangocairo). "
            "Install the system packages documented at "
            "https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation"
        ) from _WEASYPRINT_IMPORT_ERROR

    # Fetch plan data
    plan = await bcp_repo.get_plan_by_id(plan_id)
    if not plan:
        raise ValueError(f"Plan {plan_id} not found")
    
    # Gather all required data for the PDF
    data = await _gather_bcp_export_data(plan_id, event_log_limit)
    
    # Compute content hash from gathered data
    content_hash = compute_content_hash(data, {"plan_id": plan_id, "plan_title": plan["title"]})
    
    # Render HTML from BCP template
    html_content = _render_bcp_pdf_html(plan, data)
    
    # Convert HTML to PDF using WeasyPrint
    pdf_buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(pdf_buffer)
    pdf_buffer.seek(0)
    
    return pdf_buffer, content_hash


async def _gather_bcp_export_data(plan_id: int, event_log_limit: int) -> dict[str, Any]:
    """
    Gather all data needed for BCP PDF export.
    
    Args:
        plan_id: ID of the BCP plan
        event_log_limit: Maximum number of event log entries
        
    Returns:
        Dictionary containing all BCP data sections
    """
    from app.services.time_utils import humanize_hours
    
    # Section 1: Plan Overview
    objectives = await bcp_repo.list_objectives(plan_id)
    distribution_list = await bcp_repo.list_distribution_list(plan_id)
    
    # Section 2: Risk Management
    risks = await bcp_repo.list_risks(plan_id)
    insurance_policies = await bcp_repo.list_insurance_policies(plan_id)
    backup_items = await bcp_repo.list_backup_items(plan_id)
    
    # Section 3: Business Impact Analysis
    critical_activities = await bcp_repo.list_critical_activities(plan_id, sort_by="importance")
    # Add humanized RTO to each activity
    for activity in critical_activities:
        if activity.get("impact") and activity["impact"].get("rto_hours") is not None:
            activity["impact"]["rto_humanized"] = humanize_hours(activity["impact"]["rto_hours"])
    
    # Section 4: Incident Response
    checklist_items = await bcp_repo.list_checklist_items(plan_id, phase="Immediate")
    evacuation = await bcp_repo.get_evacuation_plan(plan_id)
    emergency_kit_items = await bcp_repo.list_emergency_kit_items(plan_id)
    emergency_kit_documents = [item for item in emergency_kit_items if item["category"] == "Document"]
    emergency_kit_equipment = [item for item in emergency_kit_items if item["category"] == "Equipment"]
    roles = await bcp_repo.list_roles_with_assignments(plan_id)
    contacts = await bcp_repo.list_contacts(plan_id)
    
    # Get event log - last N entries
    # First try to get from active incident, otherwise get all
    active_incident = await bcp_repo.get_active_incident(plan_id)
    if active_incident:
        all_event_log = await bcp_repo.list_event_log_entries(
            plan_id, 
            incident_id=active_incident["id"]
        )
    else:
        # Get latest entries across all incidents
        all_event_log = await bcp_repo.list_event_log_entries(plan_id)
    
    # Limit to the specified number of entries (already ordered by happened_at DESC)
    event_log = all_event_log[:event_log_limit] if len(all_event_log) > event_log_limit else all_event_log
    
    # Section 5: Recovery
    recovery_actions = await bcp_repo.list_recovery_actions(plan_id)
    # Enrich recovery actions with humanized RTO and owner names
    for action in recovery_actions:
        if action.get("rto_hours") is not None:
            action["rto_humanized"] = humanize_hours(action["rto_hours"])
        if action.get("owner_id"):
            owner = await user_repo.get_user_by_id(action["owner_id"])
            action["owner_name"] = owner.get("name") if owner else None
    
    crisis_recovery_checklist = await bcp_repo.list_checklist_items(plan_id, phase="CrisisRecovery")
    recovery_contacts = await bcp_repo.list_recovery_contacts(plan_id)
    insurance_claims = await bcp_repo.list_insurance_claims(plan_id)
    market_changes = await bcp_repo.list_market_changes(plan_id)
    
    # Section 6: Rehearse/Maintain/Review
    training_items = await bcp_repo.list_training_items(plan_id)
    review_items = await bcp_repo.list_review_items(plan_id)
    
    # Enrich roles with user names for assignments
    for role in roles:
        for assignment in role.get("assignments", []):
            if assignment.get("user_id"):
                user = await user_repo.get_user_by_id(assignment["user_id"])
                assignment["user_name"] = user.get("name") if user else "Unknown"
    
    return {
        # Section 1
        "objectives": objectives,
        "distribution_list": distribution_list,
        # Section 2
        "risks": risks,
        "insurance_policies": insurance_policies,
        "backup_items": backup_items,
        # Section 3
        "critical_activities": critical_activities,
        # Section 4
        "checklist_items": checklist_items,
        "evacuation": evacuation,
        "emergency_kit_documents": emergency_kit_documents,
        "emergency_kit_equipment": emergency_kit_equipment,
        "roles": roles,
        "contacts": contacts,
        "event_log": event_log,
        # Section 5
        "recovery_actions": recovery_actions,
        "crisis_recovery_checklist": crisis_recovery_checklist,
        "recovery_contacts": recovery_contacts,
        "insurance_claims": insurance_claims,
        "market_changes": market_changes,
        # Section 6
        "training_items": training_items,
        "review_items": review_items,
    }


def _render_bcp_pdf_html(plan: dict[str, Any], data: dict[str, Any]) -> str:
    """
    Render BCP plan to HTML using Jinja2 template.
    
    Args:
        plan: BCP plan overview data
        data: All gathered BCP data sections
        
    Returns:
        Rendered HTML string
    """
    import os
    from jinja2 import Environment, FileSystemLoader
    
    # Determine template path
    template_dir = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        "templates", "bcp", "export"
    )
    
    # Create Jinja2 environment
    env = Environment(loader=FileSystemLoader(template_dir))
    
    # Load template
    template = env.get_template("bcp_pdf.html")
    
    # Render with plan and data
    return template.render(plan=plan, **data)


def _render_plan_html(
    plan: dict[str, Any],
    version: dict[str, Any],
    metadata: dict[str, Any],
    content_json: dict[str, Any],
    template: Optional[dict[str, Any]],
    risks: list[dict[str, Any]],
) -> str:
    """Render plan content to HTML using Jinja2."""
    # HTML template with government-style formatting
    html_template = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ metadata.plan_title }}</title>
    <style>
        @page {
            size: A4;
            margin: 2cm;
            @bottom-right {
                content: "Page " counter(page) " of " counter(pages);
            }
        }
        body {
            font-family: Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }
        h1 {
            color: #003366;
            border-bottom: 3px solid #003366;
            padding-bottom: 0.5em;
            page-break-after: avoid;
        }
        h2 {
            color: #0066cc;
            margin-top: 1.5em;
            page-break-after: avoid;
        }
        h3 {
            color: #0066cc;
            margin-top: 1em;
            page-break-after: avoid;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 1em 0;
            page-break-inside: avoid;
        }
        table, th, td {
            border: 1px solid #ddd;
        }
        th {
            background-color: #003366;
            color: white;
            padding: 8px;
            text-align: left;
        }
        td {
            padding: 6px 8px;
        }
        tr:nth-child(even) {
            background-color: #f2f2f2;
        }
        .metadata-box {
            background-color: #f5f5f5;
            border: 1px solid #ddd;
            padding: 1em;
            margin: 1em 0;
        }
        .title-page {
            text-align: center;
            margin-top: 5cm;
            page-break-after: always;
        }
        .title-page h1 {
            font-size: 28pt;
            border: none;
            margin-bottom: 1em;
        }
        .section {
            page-break-inside: avoid;
        }
    </style>
</head>
<body>
    <!-- Title Page -->
    <div class="title-page">
        <h1>{{ metadata.plan_title }}</h1>
        <p style="font-size: 16pt; color: #666;">Version {{ metadata.version_number }}</p>
        {% if metadata.template_name %}
        <p style="font-size: 12pt; color: #888;">{{ metadata.template_name }}</p>
        {% endif %}
        <p style="margin-top: 3em;">
            <strong>Author:</strong> {{ metadata.author_name }}<br>
            <strong>Date:</strong> {{ metadata.authored_at | format_datetime }}
        </p>
    </div>
    
    <!-- Document Information -->
    <h1>Document Information</h1>
    <div class="metadata-box">
        <table>
            <tr>
                <th>Field</th>
                <th>Value</th>
            </tr>
            <tr>
                <td><strong>Version</strong></td>
                <td>{{ metadata.version_number }}</td>
            </tr>
            <tr>
                <td><strong>Author</strong></td>
                <td>{{ metadata.author_name }}</td>
            </tr>
            <tr>
                <td><strong>Date</strong></td>
                <td>{{ metadata.authored_at | format_datetime }}</td>
            </tr>
            {% if version.summary_change_note %}
            <tr>
                <td><strong>Change Summary</strong></td>
                <td>{{ version.summary_change_note }}</td>
            </tr>
            {% endif %}
        </table>
    </div>
    
    <!-- Plan Content -->
    {% if template and template.schema_json and template.schema_json.sections %}
        {% for section in template.schema_json.sections %}
            {% set section_id = section.section_id or section.key %}
            {% set section_content = content_json.get(section_id, {}) %}
            
            <div class="section">
                <h1>{{ section.title }}</h1>
                {% if section.description %}
                <p><em>{{ section.description }}</em></p>
                {% endif %}
                
                {% for field in section.fields %}
                    {% set field_value = section_content.get(field.field_id) %}
                    {% if field_value %}
                        <h2>{{ field.label }}</h2>
                        {% if field.field_type == 'table' and field_value is iterable %}
                            <table>
                                <thead>
                                    <tr>
                                    {% for column in field.columns %}
                                        <th>{{ column.label }}</th>
                                    {% endfor %}
                                    </tr>
                                </thead>
                                <tbody>
                                {% for row in field_value %}
                                    <tr>
                                    {% for column in field.columns %}
                                        <td>{{ row.get(column.column_id, '') }}</td>
                                    {% endfor %}
                                    </tr>
                                {% endfor %}
                                </tbody>
                            </table>
                        {% else %}
                            <p>{{ field_value }}</p>
                        {% endif %}
                    {% endif %}
                {% endfor %}
            </div>
        {% endfor %}
    {% else %}
        <!-- Fallback: render content without template structure -->
        {% for section_key, section_data in content_json.items() %}
            <h1>{{ section_key | title | replace('_', ' ') }}</h1>
            {% if section_data is mapping %}
                {% for field_key, field_value in section_data.items() %}
                    <h2>{{ field_key | title | replace('_', ' ') }}</h2>
                    <p>{{ field_value }}</p>
                {% endfor %}
            {% else %}
                <p>{{ section_data }}</p>
            {% endif %}
        {% endfor %}
    {% endif %}
    
    <!-- Appendices -->
    <h1>Appendices</h1>
    
    <!-- Risk Register -->
    {% if risks %}
    <h2>Risk Register</h2>
    <table>
        <thead>
            <tr>
                <th>Threat</th>
                <th>Likelihood</th>
                <th>Impact</th>
                <th>Rating</th>
                <th>Mitigation</th>
            </tr>
        </thead>
        <tbody>
        {% for risk in risks %}
            <tr>
                <td>{{ risk.threat }}</td>
                <td>{{ risk.likelihood or '' }}</td>
                <td>{{ risk.impact or '' }}</td>
                <td>{{ risk.rating or '' }}</td>
                <td>{{ risk.mitigation or '' }}</td>
            </tr>
        {% endfor %}
        </tbody>
    </table>
    {% endif %}
</body>
</html>
    """
    
    # Custom filter for datetime formatting
    def format_datetime(value):
        if not value:
            return ""
        if isinstance(value, str):
            try:
                dt = datetime.fromisoformat(value)
                return dt.strftime("%Y-%m-%d %H:%M UTC")
            except ValueError:
                return value
        return value
    
    # Render template using Environment with custom filter
    env = Environment(loader=BaseLoader())
    env.filters["format_datetime"] = format_datetime
    jinja_template = env.from_string(html_template)
    
    return jinja_template.render(
        plan=plan,
        version=version,
        metadata=metadata,
        content_json=content_json,
        template=template,
        risks=risks,
    )
